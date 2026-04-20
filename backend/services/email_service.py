"""
email_service.py

Handles:
- Filling Word (.docx) templates with student-specific placeholders
- Converting filled DOCX → PDF (via LibreOffice headless or fallback to reportlab)
- Sending emails via Gmail SMTP with PDF attachments
- Fetching templates from Supabase Storage
"""
import os
import io
import smtplib
import tempfile
import subprocess
from pathlib import Path
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from datetime import datetime
from typing import Optional

import httpx
from docx import Document as DocxDocument
from dotenv import load_dotenv
from sqlalchemy.orm import Session

from models.db_models import (
    Application, User, Course, ExamSchedule,
    Template, EmailLog, EmailStatus, EmailType, ApplicationStatus, TemplateType,
    Document, DocumentType
)

load_dotenv(dotenv_path=Path(__file__).resolve().parent.parent / ".env")

GMAIL_SENDER       = os.getenv("GMAIL_SENDER")
GMAIL_APP_PASSWORD = os.getenv("GMAIL_APP_PASSWORD")
SUPABASE_URL       = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY")

COLLEGE_NAME    = "Institute of Engineering and Management, Kolkata"
COLLEGE_ADDRESS = "GN-34/2, Sector-V, Salt Lake Electronics Complex, Kolkata - 700091"
COLLEGE_EMAIL   = "iemadmission2@gmail.com"
COLLEGE_PHONE   = "+91 33 2357 2995"


# ── Fetch template file from Supabase Storage ──────────────

async def fetch_template_bytes(file_url: str) -> bytes:
    async with httpx.AsyncClient(timeout=30.0) as client:
        resp = await client.get(
            file_url,
            headers={"Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"}
        )
        if resp.status_code != 200:
            raise Exception(f"Failed to fetch template: {resp.text}")
        return resp.content


# ── Build placeholder dict for a student ──────────────────

def build_placeholders(
    student: User,
    application: Application,
    course: Course,
    schedule: Optional[ExamSchedule] = None,
    roll_number: Optional[str] = None
) -> dict:
    """
    All {{placeholders}} that can appear in director-uploaded templates.
    Director uses these exact keys in their Word template.
    """
    today = datetime.now().strftime("%d %B %Y")

    placeholders = {
        "{{student_name}}":     student.name,
        "{{student_email}}":    student.email,
        "{{course_name}}":      course.name,
        "{{course_type}}":      course.type.value if hasattr(course.type, 'value') else str(course.type),
        "{{application_id}}":   str(application.id),
        "{{college_name}}":     COLLEGE_NAME,
        "{{college_address}}":  COLLEGE_ADDRESS,
        "{{college_email}}":    COLLEGE_EMAIL,
        "{{college_phone}}":    COLLEGE_PHONE,
        "{{date}}":             today,
        "{{fees}}":             f"Rs. {course.fees:,.0f}",
        "{{status}}":           application.status,
    }

    if schedule:
        placeholders.update({
            "{{exam_date}}":   schedule.exam_date.strftime("%d %B %Y, %I:%M %p") if schedule.exam_date else "TBA",
            "{{venue}}":       schedule.venue or "TBA",
            "{{result_date}}": schedule.result_release_date.strftime("%d %B %Y") if schedule.result_release_date else "TBA",
            "{{syllabus_url}}":schedule.syllabus_url or "Available on college website",
        })

    if roll_number:
        placeholders["{{roll_number}}"] = roll_number

    return placeholders


# ── Fill DOCX template with placeholders ──────────────────

def fill_docx_template(template_bytes: bytes, placeholders: dict) -> bytes:
    """Replace all {{placeholder}} strings in a DOCX template"""
    doc = DocxDocument(io.BytesIO(template_bytes))

    def replace_in_paragraph(paragraph):
        # 1. Try simple run-level replacement (preserves formatting)
        for run in paragraph.runs:
            for key, value in placeholders.items():
                if key in run.text:
                    run.text = run.text.replace(key, str(value))

        # 2. Check if any placeholders are still present (split across runs)
        for key, value in placeholders.items():
            if key in paragraph.text:
                # Fallback: Merge runs (destroys fine-grained formatting but fixes the placeholder)
                full_text = "".join(run.text for run in paragraph.runs).replace(key, str(value))
                if paragraph.runs:
                    paragraph.runs[0].text = full_text
                    for run in paragraph.runs[1:]:
                        run.text = ""

    for para in doc.paragraphs:
        replace_in_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


# ── Convert DOCX → PDF ────────────────────────────────────

def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """
    Uses LibreOffice headless for conversion (available on Linux/Render).
    Falls back to a simple reportlab PDF if LibreOffice not available.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / "document.docx"
        pdf_path  = Path(tmpdir) / "document.pdf"

        docx_path.write_bytes(docx_bytes)

        try:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf",
                 "--outdir", tmpdir, str(docx_path)],
                check=True,
                timeout=30,
                capture_output=True
            )
            return pdf_path.read_bytes()

        except (subprocess.CalledProcessError, FileNotFoundError):
            # Fallback: basic PDF using reportlab
            print("[EmailService] WARNING: LibreOffice not found. Falling back to basic PDF (formatting/logos will be lost).")
            return _fallback_pdf(docx_bytes)


def _fallback_pdf(docx_bytes: bytes) -> bytes:
    """Minimal PDF fallback using reportlab if LibreOffice unavailable (dev mode)"""
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet

    # Extract text from DOCX for the fallback
    doc  = DocxDocument(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    buffer = io.BytesIO()
    pdf    = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story  = []

    for line in text.split("\n"):
        if line.strip():
            story.append(Paragraph(line, styles["Normal"]))
            story.append(Spacer(1, 6))

    pdf.build(story)
    return buffer.getvalue()


# ── Generate roll number ───────────────────────────────────

def generate_roll_number(application_id: int, course: Course) -> str:
    year   = datetime.now().year
    prefix = "".join(w[0] for w in course.name.split()[:2]).upper()
    return f"IEM/{year}/{prefix}/{application_id:04d}"


# ── Core email sender ──────────────────────────────────────

def send_email(
    to_email: str,
    subject: str,
    body_html: str,
    attachments: list[tuple[str, bytes]] = []   # [(filename, bytes)]
) -> bool:
    """
    Sends an email via Gmail SMTP.
    Returns True on success, False on failure.
    """
    try:
        msg = MIMEMultipart("mixed")
        msg["From"]    = GMAIL_SENDER
        msg["To"]      = to_email
        msg["Subject"] = subject

        msg.attach(MIMEText(body_html, "html"))

        for filename, file_bytes in attachments:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(file_bytes)
            encoders.encode_base64(part)
            part.add_header("Content-Disposition", f'attachment; filename="{filename}"')
            msg.attach(part)

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
            smtp.login(GMAIL_SENDER, GMAIL_APP_PASSWORD)
            smtp.sendmail(GMAIL_SENDER, to_email, msg.as_string())

        return True

    except Exception as e:
        print(f"[EmailService] Failed to send to {to_email}: {e}")
        return False


# ── Result email (selected) ────────────────────────────────

async def send_selected_email(
    db: Session,
    application: Application,
    student: User,
    course: Course,
    schedule: Optional[ExamSchedule]
) -> bool:
    roll_no      = generate_roll_number(application.id, course)
    placeholders = build_placeholders(student, application, course, schedule, roll_no)

    # Try to use director's uploaded template
    template = db.query(Template).filter(
        Template.type == TemplateType.result_selected
    ).first()

    attachments = []

    if template:
        try:
            template_bytes = await fetch_template_bytes(template.file_url)
            filled_docx    = fill_docx_template(template_bytes, placeholders)
            pdf_bytes      = convert_docx_to_pdf(filled_docx)
            attachments.append(("Selection_Letter.pdf", pdf_bytes))
        except Exception as e:
            print(f"[EmailService] Template error: {e}, using fallback email body")

    # Attach admit card if template exists
    admit_template = db.query(Template).filter(
        Template.type == TemplateType.admit_card
    ).first()

    if admit_template:
        try:
            admit_bytes  = await fetch_template_bytes(admit_template.file_url)
            filled_admit = fill_docx_template(admit_bytes, placeholders)
            admit_pdf    = convert_docx_to_pdf(filled_admit)
            attachments.append(("Admit_Card.pdf", admit_pdf))
        except Exception as e:
            print(f"[EmailService] Admit card error: {e}")

    # Attach syllabus PDF if uploaded
    if schedule and schedule.syllabus_url:
        try:
            async with httpx.AsyncClient() as client:
                syllabus_resp = await client.get(schedule.syllabus_url)
            if syllabus_resp.status_code == 200:
                attachments.append(("Syllabus.pdf", syllabus_resp.content))
        except Exception as e:
            print(f"[EmailService] Syllabus fetch error: {e}")

    subject   = f"🎉 Congratulations! You are Selected – {course.name} | {COLLEGE_NAME}"
    body_html = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: auto;">
        <h2 style="color: #1a5276;">Congratulations, {student.name}!</h2>
        <p>We are delighted to inform you that you have been <strong>selected</strong>
        for <strong>{course.name}</strong> at {COLLEGE_NAME}.</p>

        <table style="border-collapse: collapse; width: 100%; margin: 20px 0;">
            <tr><td style="padding: 8px; font-weight: bold;">Roll Number</td>
                <td style="padding: 8px;">{roll_no}</td></tr>
            <tr style="background:#f2f2f2;"><td style="padding: 8px; font-weight: bold;">Course</td>
                <td style="padding: 8px;">{course.name}</td></tr>
            <tr><td style="padding: 8px; font-weight: bold;">Exam Date</td>
                <td style="padding: 8px;">{placeholders.get('{{exam_date}}', 'TBA')}</td></tr>
            <tr style="background:#f2f2f2;"><td style="padding: 8px; font-weight: bold;">Venue</td>
                <td style="padding: 8px;">{placeholders.get('{{venue}}', 'TBA')}</td></tr>
        </table>

        <p>Please find your <strong>Selection Letter</strong> and <strong>Admit Card</strong> attached.</p>
        <p>For queries, contact us at <a href="mailto:{COLLEGE_EMAIL}">{COLLEGE_EMAIL}</a>
        or call {COLLEGE_PHONE}.</p>

        <hr/>
        <p style="font-size: 12px; color: #888;">
            {COLLEGE_NAME}<br/>{COLLEGE_ADDRESS}
        </p>
    </div>
    """

    success = send_email(student.email, subject, body_html, attachments)
    _log_email(db, student.id, EmailType.result, success, result_status=application.status)
    return success


# ── Result email (rejected) ────────────────────────────────

async def send_rejected_email(
    db: Session,
    application: Application,
    student: User,
    course: Course
) -> bool:
    placeholders = build_placeholders(student, application, course)
    attachments  = []

    template = db.query(Template).filter(
        Template.type == TemplateType.result_rejected
    ).first()

    if template:
        try:
            template_bytes = await fetch_template_bytes(template.file_url)
            filled_docx    = fill_docx_template(template_bytes, placeholders)
            pdf_bytes      = convert_docx_to_pdf(filled_docx)
            attachments.append(("Result_Letter.pdf", pdf_bytes))
        except Exception as e:
            print(f"[EmailService] Rejected template error: {e}")

    subject   = f"Admission Update – {course.name} | {COLLEGE_NAME}"
    body_html = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: auto;">
        <h2 style="color: #922b21;">Dear {student.name},</h2>
        <p>Thank you for your interest in <strong>{course.name}</strong>
        at {COLLEGE_NAME}.</p>
        <p>After careful review of your application, we regret to inform you that
        we are unable to offer you admission in this cycle.</p>
        <p>We encourage you to apply again in the next session.
        For feedback or queries, please contact
        <a href="mailto:{COLLEGE_EMAIL}">{COLLEGE_EMAIL}</a>.</p>
        <p>We wish you the very best in your academic journey.</p>
        <hr/>
        <p style="font-size: 12px; color: #888;">
            {COLLEGE_NAME}<br/>{COLLEGE_ADDRESS}
        </p>
    </div>
    """

    success = send_email(student.email, subject, body_html, attachments)
    _log_email(db, student.id, EmailType.result, success)
    return success


# ── Admit Card Storage ─────────────────────────────────────

async def generate_and_store_admit_card(db: Session, application: Application) -> str:
    """
    Core logic to generate an admit card PDF and store it in Supabase + Document table.
    Returns the public URL of the generated PDF.
    """
    student  = db.query(User).filter(User.id == application.student_id).first()
    course   = db.query(Course).filter(Course.id == application.course_id).first()
    schedule = db.query(ExamSchedule).filter(ExamSchedule.course_id == application.course_id).first()

    # 1. Fetch template
    template = db.query(Template).filter(Template.type == TemplateType.admit_card).first()
    if not template:
        raise ValueError("Admit card template not found")

    template_bytes = await fetch_template_bytes(template.file_url)

    # 2. Generate PDF
    roll_no      = generate_roll_number(application.id, course)
    placeholders = build_placeholders(student, application, course, schedule, roll_no)
    filled_docx  = fill_docx_template(template_bytes, placeholders)
    pdf_bytes    = convert_docx_to_pdf(filled_docx)

    # 3. Upload to Supabase
    storage_filename = f"{student.id}_admit_card.pdf"
    upload_url       = f"{SUPABASE_URL}/storage/v1/object/admit-cards/{storage_filename}"

    async with httpx.AsyncClient(timeout=60.0) as client:
        resp = await client.post(
            upload_url,
            content=pdf_bytes,
            headers={
                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                "Content-Type":  "application/pdf",
                "x-upsert":      "true"
            }
        )
        if resp.status_code not in (200, 201):
            raise Exception(f"Supabase upload failed: {resp.text}")

    file_url = f"{SUPABASE_URL}/storage/v1/object/public/admit-cards/{storage_filename}"

    # 4. Save/Update Document record
    existing = db.query(Document).filter(
        Document.student_id == student.id,
        Document.type       == DocumentType.admit_card
    ).first()

    if existing:
        existing.file_url = file_url
    else:
        db.add(Document(
            student_id      = student.id,
            type            = DocumentType.admit_card,
            file_url        = file_url,
            extracted_marks = None
        ))
    
    db.commit()
    return file_url


# ── Email log helper ───────────────────────────────────────

def _log_email(db: Session, student_id: int, email_type: EmailType, success: bool, error: str = None, result_status: str = None):
    log = EmailLog(
        student_id = student_id,
        type       = email_type,
        result_status = result_status,
        status     = EmailStatus.sent if success else EmailStatus.failed,
        error_msg  = error if not success else None
    )
    db.add(log)
    db.commit()